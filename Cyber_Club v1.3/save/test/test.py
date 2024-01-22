import docx # Подключение к docx
from docx.shared import Pt # Подключение размер шрифта


doc = docx.Document()

style = doc.styles['Normal'] # Стили текста
style.font.name = 'Courier New CYR' # Шрифт
style.font.size = Pt(22) # Размер шрифта

doc.add_paragraph('Абзац!')

par1 = doc.add_paragraph('')
par2 = doc.add_paragraph('Второй.')
par3 = doc.add_paragraph('Третий.')

# Дополнение абзацов

par1.add_run('Продолжение Первый абзаца.').italic = True # Курсив
par2.add_run('Продолжение Второй абзаца.').bold = True # Жирный

run = par3.add_run('Продолжение Третий абзаца.')
run.bold = True # Жирный
run.underline = True # Линий

doc.save('test.docx') # Сохранение/скачивание