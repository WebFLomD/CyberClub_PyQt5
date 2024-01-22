import docx # Подключение к docx
from docx.shared import Pt # Подключение размер шрифта


doc = docx.Document()

style = doc.styles['Normal'] # Стили текста
style.font.name = 'Courier New CYR' # Шрифт
style.font.size = Pt(22) # Размер шрифта

doc.add_paragraph('ООО "Cyber Club"')

par1 = doc.add_paragraph('')
par2 = doc.add_paragraph('')
par3 = doc.add_paragraph('')
par4 = doc.add_paragraph('')
par5 = doc.add_paragraph('')
par6 = doc.add_paragraph('')
par7 = doc.add_paragraph('')
par8 = doc.add_paragraph('')
par9 = doc.add_paragraph('')
par10 = doc.add_paragraph('')
par11 = doc.add_paragraph('')
par12 = doc.add_paragraph('')
par13 = doc.add_paragraph('')
par14 = doc.add_paragraph('')


#######

par1.add_run('Добро пожаловать')
par2.add_run('ККМ 00000001     #0001')
par3.add_run('ИНН 0000000000001')
par4.add_run('      ЭКЛЗ 0000000001')
par5.add_run('11.11.11 11:11 СИС.')
par6.add_run('АДМИН - TEST')
par7.add_run('АРЕНДА ПК НА 1 час')
par8.add_run('=75.00')
par9.add_run('______________________')
par10.add_run('ИТОГ').bold=True
par11.add_run('     75.00').bold=True
par12.add_run('НАЛИЧ/КАРТОЙ     =75.00')
par13.add_run('***********************')
par14.add_run('       00000001# 000001')

doc.save('test.docx') # Сохранение/скачивание